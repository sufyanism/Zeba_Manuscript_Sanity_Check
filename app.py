import streamlit as st
import re
from io import BytesIO
from docx import Document

def process_manuscript(file):
    doc = Document(file)
    report = []
    space_count = 0
    quote_count = 0
    
    # Define Unicode typographic characters outside of regex to avoid "bad escape" errors
    lsquo, rsquo = '\u2018', '\u2019'  # Single open, close/apostrophe
    ldquo, rdquo = '\u201c', '\u201d'  # Double open, close

    # --- 1. HEADING AUDIT ---
    # Looks for 'Heading 1' regardless of case
    headings = [p.text for p in doc.paragraphs if "heading 1" in p.style.name.lower()]
    
    report.append("### 📂 Heading Audit")
    if not headings:
        report.append("⚠️ **Warning:** No 'Heading 1' styles found. Please check your Word styles.")
    else:
        for i, text in enumerate(headings, 1):
            nums = re.findall(r'\d+', text)
            if nums:
                actual_num = int(nums[0])
                if actual_num != i:
                    report.append(f"⚠️ **Sequence Break:** Found 'Chapter {actual_num}', but expected position {i}.")
            else:
                report.append(f"ℹ️ **Note:** Heading found without a number: '{text[:30]}...'")

    # --- 2. TYPOGRAPHIC CLEANUP ---
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
            
        initial_text = para.text
        
        # A. Collapse Multiple Spaces
        if re.search(r' {2,}', para.text):
            orig_len = len(para.text)
            para.text = re.sub(r' {2,}', ' ', para.text)
            space_count += (orig_len - len(para.text))

        # B. Smart Quote Logic (Concatenating characters to avoid \u escape errors)
        # Double Quotes: Open if start of string or following whitespace/bracket
        para.text = re.sub(r'(^|[\s(\[])"', r'\1' + ldquo, para.text) 
        para.text = re.sub(r'"', rdquo, para.text) # Close all others
        
        # Single Quotes: Open if start of string or following whitespace/bracket
        para.text = re.sub(r"(^|[\s(\[])'", r"\1" + lsquo, para.text) 
        para.text = re.sub(r"'", rsquo, para.text) # Close/Apostrophe all others
        
        if initial_text != para.text:
            quote_count += 1

    report.append("### 🛠️ Cleanup Actions")
    report.append(f"✅ **Spaces:** Removed {space_count} redundant spaces.")
    report.append(f"✅ **Typography:** Updated quotes/apostrophes in {quote_count} paragraphs.")
    
    # Save to memory buffer
    out_io = BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    
    return out_io, "\n\n".join(report)

# --- STREAMLIT UI ---
st.set_page_config(page_title="The Pickaxe", page_icon="⛏️")

st.title("⛏️ The Pickaxe")
st.write("Professional Manuscript Auditor & Typographic Cleaner")

uploaded_file = st.file_uploader("Upload Manuscript (.docx)", type="docx")

if uploaded_file:
    with st.spinner("Analyzing manuscript..."):
        try:
            cleaned_file, sanity_report = process_manuscript(uploaded_file)
            
            st.success("Processing Complete!")
            
            with st.expander("View Sanity Report", expanded=True):
                st.markdown(sanity_report)
            
            st.download_button(
                label="📥 Download Cleaned Document",
                data=cleaned_file,
                file_name=f"CLEANED_{uploaded_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Critical Error: {e}")
            st.info("Check if the document is password protected or corrupted.")